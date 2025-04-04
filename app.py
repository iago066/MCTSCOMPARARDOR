# App Tkinter com comparação e exportação de publicações únicas
# Desenvolvido por MARIA CLARA NOGUEIRA DINIZ OAB/PI 23.765

import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from docx import Document
from docx.shared import RGBColor

import re
import docx2txt
import os

from text_extractor import TextExtractor
from publication_extractor import PublicationExtractor

class ComparadorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Comparador de Publicações Jurídicas")
        self.text_extractor = TextExtractor()
        self.publication_extractor = PublicationExtractor()
        self.files = []
        self.setup_ui()

    def setup_ui(self):
        self.root.geometry("600x350")
        self.root.configure(bg='#f8f9fa')
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TFrame', background='#f8f9fa')
        style.configure('TLabel', background='#f8f9fa', font=('Segoe UI', 10))
        style.configure('TButton', font=('Segoe UI', 10), padding=6)
        rodape_label = ttk.Label(self.root, text="Desenvolvido por MARIA CLARA NOGUEIRA DINIZ OAB/PI 23.765", foreground="#888888", font=('Segoe UI', 9))
        rodape_label.pack(side=tk.BOTTOM, pady=5)
        frm = ttk.Frame(self.root, padding=10)
        frm.pack()

        ttk.Label(frm, text="Selecione até 3 arquivos:").grid(row=0, column=0, columnspan=3)

        self.entries = []
        for i in range(3):
            entry = ttk.Entry(frm, width=60)
            entry.grid(row=i+1, column=0)
            self.entries.append(entry)
            ttk.Button(frm, text="Selecionar", command=lambda e=entry: self.selecionar_arquivo(e)).grid(row=i+1, column=1)

        ttk.Button(frm, text="Comparar e Exportar", command=self.comparar).grid(row=4, column=0, columnspan=2, pady=10)

    def selecionar_arquivo(self, entry):
        path = filedialog.askopenfilename(title="Selecionar documento .docx")
        if path:
            entry.delete(0, tk.END)
            entry.insert(0, path)

    def comparar(self):
        paths = [e.get() for e in self.entries if e.get()]
        if len(paths) < 2:
            messagebox.showerror("Erro", "Selecione pelo menos 2 arquivos.")
            return

        dados = {}
        for path in paths:
            nome = os.path.basename(path)
            texto = self.text_extractor.extract_text(path)
            processos = list(set(re.findall(self.publication_extractor.case_number_pattern, texto)))
            pubs = self.publication_extractor.extract_publications(texto, processos)
            dados[nome] = {"processos": set(pubs.keys()), "publicacoes": pubs, "total": len(pubs)}

        nomes = list(dados.keys())
        combinados = []
        idx = 0

        for i, nome in enumerate(nomes):
            unicos = dados[nome]["processos"]
            for j, outro in enumerate(nomes):
                if i != j:
                    unicos -= dados[outro]["processos"]
            for proc in sorted(unicos):
                texto = dados[nome]["publicacoes"][proc]["texto_completo"]
                match = re.search(r'Publica[çc][aã]o[:\s]*\d+\s+de\s+\d+', texto, re.IGNORECASE)
                ident = match.group(0) if match else f"Publicação gerada nº {str(idx+1).zfill(3)}"
                combinados.append({
                    "Número do Processo": proc,
                    "Fonte": nome,
                    "Publicação Original": ident,
                    "Texto Completo": texto
                })
                idx += 1

        doc = Document()
        doc.add_heading("Diário Consolidado - Comparação de Publicações", level=1)
        for nome in nomes:
            doc.add_paragraph(f"Total de publicações em '{nome}': {dados[nome]['total']}")
        doc.add_paragraph(f"Total de publicações únicas: {len(combinados)}\n")

        cores = [RGBColor(255,0,0), RGBColor(0,0,255), RGBColor(0,0,0)]
        cor_mapa = {n: cores[i] for i, n in enumerate(nomes)}

        for pub in combinados:
            cor = cor_mapa.get(pub["Fonte"], RGBColor(0,0,0))
            titulo = doc.add_paragraph()
            run = titulo.add_run(pub["Publicação Original"])
            run.bold = True
            run.font.color.rgb = cor

            p1 = doc.add_paragraph()
            run = p1.add_run(f"Número do Processo: {pub['Número do Processo']}")
            run.font.color.rgb = cor

            p2 = doc.add_paragraph()
            run = p2.add_run(f"Fonte: {pub['Fonte']}")
            run.font.color.rgb = cor

            p3 = doc.add_paragraph()
            run = p3.add_run(pub["Texto Completo"])
            run.font.color.rgb = cor

            doc.add_paragraph("\n" + "-" * 80 + "\n")

        out_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if out_path:
            doc.save(out_path)
            messagebox.showinfo("Sucesso", f"Diário exportado para:\n{out_path}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ComparadorApp(root)
    root.mainloop()